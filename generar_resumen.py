"""
Genera un resumen en Word (.docx) del avance del proyecto en una fecha dada.
No forma parte del chatbot: es solo una utilidad para reportar avance.

Uso:
    .venv\\Scripts\\python generar_resumen.py
"""

from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HOY = date(2026, 7, 10)
NOMBRE_ARCHIVO = f"Resumen_avance_{HOY.isoformat()}.docx"


def agregar_titulo(doc, texto, nivel=1):
    doc.add_heading(texto, level=nivel)


def agregar_lista(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    # Estilo base
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    # --- Portada simple ---
    titulo = doc.add_heading("Chatbot de Autodiagnóstico", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Resumen de avance del proyecto")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    fecha_p = doc.add_paragraph(HOY.strftime("%d de %B de %Y"))
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- 1. Contexto ---
    agregar_titulo(doc, "1. ¿Qué se está construyendo?")
    doc.add_paragraph(
        "Un chatbot con IA para uso interno de la compañía, que responde en "
        "lenguaje natural preguntas sobre el proceso de Autodiagnóstico "
        "(diagnóstico del módem de wifi del cliente), consultando datos reales "
        "de los tres canales del proceso: Portal Web, Bot de WhatsApp y "
        "Sysbrazo (CRM de asesores)."
    )

    # --- 2. Fase 1 ---
    agregar_titulo(doc, "2. Fase 1 — Prototipo funcionando (COMPLETADA)")
    doc.add_paragraph(
        "Se construyó una página web de chat funcionando localmente, con datos "
        "de ejemplo (simulados pero realistas), para validar la idea antes de "
        "conectar las bases de datos reales."
    )
    agregar_lista(doc, [
        "Página de chat en Streamlit: se escribe una pregunta y responde con "
        "cifras y un gráfico.",
        "Las 10 mediciones definidas para el proyecto, cada una calculada y "
        "verificada sobre datos de ejemplo (~4.000 autodiagnósticos, 8 ciudades, "
        "3 canales, 3 meses de historia).",
        "Filtros por ciudad, fecha y canal en la barra lateral.",
        "El chatbot funciona con o sin conexión a la IA de Claude: sin ella "
        "reconoce preguntas directas; con ella, entiende preguntas escritas "
        "libremente.",
        "Servidor probado: arranca sin errores y responde correctamente.",
    ])

    doc.add_paragraph(
        "→ Aquí se puede insertar el pantallazo del chatbot funcionando."
    )
    doc.add_paragraph("[ ESPACIO PARA PANTALLAZO DEL CHATBOT ]").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()

    # --- 3. Fase 2 ---
    agregar_titulo(doc, "3. Fase 2 — Conexión a datos reales (EN PROGRESO)")
    doc.add_paragraph(
        "Esta fase reemplaza los datos de ejemplo por las bases de datos reales "
        "de la compañía: sysbrazo, analytics y odoo (Postgres)."
    )

    agregar_titulo(doc, "3.1 Avance de hoy", nivel=2)
    agregar_lista(doc, [
        "Se definió el método de conexión: conexión directa a la base de datos "
        "Postgres (en vez de vía Redash).",
        "Se documentaron y confirmaron con datos reales los campos clave de "
        "Odoo (helpdesk_ticket): estado del ticket (Solved = resuelto), equipo "
        "que atiende (NOC / OPS / Otro), tiempo de resolución (close_hours) y "
        "motivo de cierre (fbz_helpdesk_motivo).",
        "Se confirmó el campo que conecta un autodiagnóstico de Sysbrazo con su "
        "ticket en Odoo (ticket_ref).",
        "Se resolvieron 3 definiciones de negocio pendientes: el equipo "
        "'NET Operations' es distinto a NOC; el motivo de cierre sale de "
        "fbz_helpdesk_motivo; y 'terminó resuelto' incluye tanto autodiagnósticos "
        "completados como tickets que quedaron Solved en Odoo.",
        "Se redactaron las 10 consultas SQL (una por cada medición del "
        "proyecto), listas para probarse contra la base de datos real.",
    ])

    agregar_titulo(doc, "3.2 Qué falta para completar la Fase 2", nivel=2)
    agregar_lista(doc, [
        "Probar las 10 consultas contra la base de datos real y validar que "
        "las cifras resultantes tienen sentido.",
        "Obtener las credenciales de conexión a Postgres (de forma segura, sin "
        "compartirlas por chat).",
        "Conectar esas consultas al chatbot, reemplazando los datos de ejemplo.",
    ])

    # --- 4. Próximas fases ---
    agregar_titulo(doc, "4. Próximas fases")
    agregar_lista(doc, [
        "Fase 3: afinar las 10 mediciones ya con datos reales.",
        "Fase 4: sumar los canales Bot y Portal (datos en Mixpanel).",
        "Fase 5: publicar el chatbot para que otras personas de la compañía "
        "puedan usarlo.",
    ])

    doc.save(NOMBRE_ARCHIVO)
    print(f"Documento generado: {NOMBRE_ARCHIVO}")


if __name__ == "__main__":
    main()
